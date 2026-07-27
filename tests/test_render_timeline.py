import tempfile
import unittest
from pathlib import Path

from docx import Document

import render_timeline as timeline


def make_doc(path: Path, headers, rows):
    doc = Document()
    doc.add_heading("EVENTS", 1)
    table = doc.add_table(rows=1, cols=len(headers))
    for i, value in enumerate(headers):
        table.rows[0].cells[i].text = value
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    doc.save(path)


class TimelineTests(unittest.TestCase):
    def test_aliases_and_multiple_date_formats(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "input.docx"
            make_doc(path, ["Event", "When", "Finish", "Theme"], [
                ["First", "27/07/2026", "", "A"],
                ["Second", "1 August 2026", "5 August 2026", "B"],
            ])
            result = timeline.read_docx(path)
            self.assertEqual(2, len(result.events))
            self.assertEqual("2026-07-27", result.events[0]["start"])
            self.assertEqual("2026-08-05", result.events[1]["end"])
            self.assertFalse([x for x in result.issues if x.level == "error"])

    def test_invalid_rows_are_reported_and_skipped(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "input.docx"
            make_doc(path, ["Date", "Title"], [
                ["not-a-date", "Bad"],
                ["2026-07-27", "Good"],
                ["2026-07-28", ""],
            ])
            result = timeline.read_docx(path)
            self.assertEqual(["Good"], [x["title"] for x in result.events])
            self.assertEqual(2, len([x for x in result.issues if x.level == "error"]))

    def test_duplicate_ids_are_made_unique(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "input.docx"
            make_doc(path, ["ID", "Start", "Title"], [
                ["same", "2026-01-01", "One"],
                ["same", "2026-01-02", "Two"],
            ])
            result = timeline.read_docx(path)
            self.assertEqual(["same", "same-2"], [x["id"] for x in result.events])

    def test_unsafe_link_is_removed(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "input.docx"
            make_doc(path, ["Date", "Title", "Link"], [
                ["2026-01-01", "One", "javascript:alert(1)"],
            ])
            result = timeline.read_docx(path)
            self.assertEqual("", result.events[0]["link"])
            self.assertTrue(any("link" in x.message for x in result.issues))


if __name__ == "__main__":
    unittest.main()
