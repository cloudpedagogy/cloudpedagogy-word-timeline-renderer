#!/usr/bin/env python3
"""Create the documented sample Word input for render_timeline.py."""

from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "examples" / "timeline_example.docx"


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(document, headers, rows, widths=None, font_size=9):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_header(table.rows[0])
    for i, value in enumerate(headers):
        table.rows[0].cells[i].text = value
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = str(value)
    for row in table.rows:
        prevent_row_split(row)
        for i, cell in enumerate(row.cells):
            if widths:
                cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
    return table


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    title = doc.add_heading("Public Health Programme Timeline", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        "Edit the SETTINGS and EVENTS tables, save the document, then run render_timeline.py. "
        "Column names are case-insensitive; several common alternatives are accepted."
    )
    p.style = doc.styles["Intense Quote"]

    doc.add_heading("SETTINGS", level=1)
    add_table(doc, ["Setting", "Value"], [
        ["Title", "Public Health Programme Redesign"],
        ["Subtitle", "Illustrative milestones, workstreams and approval stages"],
        ["Height", "620px"],
        ["Orientation", "bottom"],
        ["Stack", "yes"],
    ], widths=[1.8, 7.8], font_size=9.5)

    doc.add_heading("EVENTS", level=1)
    add_table(doc,
        ["Event ID", "Start Date", "End Date", "Title", "Description", "Category", "Track", "Colour"],
        [
            ["e01", "2026-01-12", "2026-02-20", "Discovery and evidence review",
             "Review the existing programme, student feedback and assessment pattern.",
             "Planning", "Programme", "#bfdbfe"],
            ["e02", "2026-02-02", "2026-03-13", "Module mapping workshops",
             "Map learning outcomes, content and assessments across modules.",
             "Curriculum", "Academic design", "#bbf7d0"],
            ["e03", "2026-03-02", "2026-04-03", "Assessment redesign",
             "Develop a smaller, coherent programme-level assessment pattern.",
             "Assessment", "Academic design", "#fde68a"],
            ["e04", "2026-04-15", "", "Stage 3 review",
             "Submit the redesign for the third programme approval stage.",
             "Approval", "Governance", "#fecaca"],
            ["e05", "2026-04-20", "2026-05-29", "Learning design prototypes",
             "Prototype online and intensive-mode learning activities.",
             "Development", "Digital education", "#ddd6fe"],
            ["e06", "2026-06-17", "", "Stage 4 approval",
             "Final institutional approval milestone.",
             "Approval", "Governance", "#fecaca"],
            ["e07", "2026-07-01", "2027-05-31", "Content production",
             "Produce, review and quality-assure module materials.",
             "Development", "Digital education", "#ddd6fe"],
            ["e08", "2027-09-27", "", "Revised programme launches",
             "First teaching week for the redesigned programme.",
             "Delivery", "Programme", "#a7f3d0"],
        ], widths=[0.65, 0.9, 0.9, 1.45, 2.65, 1.0, 1.25, 0.8], font_size=8.5)

    doc.add_heading("Column guidance", level=1)
    guidance = [
        ("Event ID", "Optional unique identifier. A value is generated if blank."),
        ("Start Date", "Required. Use YYYY-MM-DD, DD/MM/YYYY, or a written date such as 15 April 2026."),
        ("End Date", "Optional. When supplied, the item is displayed as a date range."),
        ("Title", "Required event label."),
        ("Description", "Optional text displayed in the event tooltip."),
        ("Category", "Optional filter category."),
        ("Track", "Optional timeline lane/group."),
        ("Colour", "Optional CSS colour name, hex, rgb() or hsl()."),
        ("URL", "Optional http/https link."),
    ]
    for name, text in guidance:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(f"{name}: ")
        run.bold = True
        para.add_run(text)
    for style_name in ("Normal", "Body Text"):
        if style_name in doc.styles:
            doc.styles[style_name].font.name = "Arial"
            doc.styles[style_name].font.size = Pt(10.5)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
